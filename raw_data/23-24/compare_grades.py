#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
比对PDF和Word文档中的学号和智育成绩
"""

import sys
import pandas as pd
from pathlib import Path
import warnings
warnings.filterwarnings('ignore')

try:
    import pdfplumber
except ImportError:
    print("请安装pdfplumber: pip install pdfplumber")
    sys.exit(1)

try:
    from docx import Document
except ImportError:
    print("请安装python-docx: pip install python-docx")
    sys.exit(1)

def extract_pdf_data(pdf_path):
    """从PDF文件提取学号和智育成绩"""
    print(f"正在读取PDF文件: {pdf_path}")
    
    data = []
    with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            print(f"处理PDF第{page_num + 1}页")
            
            # 尝试提取表格
            tables = page.extract_tables()
            if tables:
                for table in tables:
                    for row in table:
                        if row and len(row) >= 4:  # 至少需要4列
                            # 跳过表头行
                            if row[0] and str(row[0]).strip() and not str(row[0]).strip().startswith(('序号', '排名', '姓名')):
                                student_id = str(row[0]).strip()
                                # 第四列是智育成绩
                                if len(row) > 3:
                                    grade = str(row[3]).strip() if row[3] else ""
                                    data.append({
                                        '学号': student_id,
                                        '智育成绩': grade,
                                        '页面': page_num + 1
                                    })
            
            # 如果没有找到表格，尝试文本提取
            if not tables:
                text = page.extract_text()
                if text:
                    lines = text.split('\n')
                    for line in lines:
                        # 简单的正则匹配学号模式 (通常是数字)
                        import re
                        parts = re.split(r'\s+', line.strip())
                        if len(parts) >= 4 and parts[0].isdigit():
                            data.append({
                                '学号': parts[0],
                                '智育成绩': parts[3],
                                '页面': page_num + 1
                            })
    
    print(f"从PDF提取了 {len(data)} 条记录")
    return data

def extract_word_data(docx_path):
    """从Word文档提取学号和智育成绩"""
    print(f"正在读取Word文档: {docx_path}")
    
    data = []
    doc = Document(docx_path)
    
    # 从表格中提取数据
    for table_idx, table in enumerate(doc.tables):
        print(f"处理Word表格 {table_idx + 1}")
        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 4:
                # 跳过表头
                if cells[0] and not cells[0].startswith(('序号', '排名', '姓名')):
                    student_id = cells[0]
                    grade = cells[3] if len(cells) > 3 else ""
                    
                    data.append({
                        '学号': student_id,
                        '智育成绩': grade,
                        '表格': table_idx + 1,
                        '行': row_idx + 1
                    })
    
    # 如果没有表格或表格为空，尝试从段落中提取
    if not data:
        for para_idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                import re
                parts = re.split(r'\s+', text)
                if len(parts) >= 4 and parts[0].isdigit():
                    data.append({
                        '学号': parts[0],
                        '智育成绩': parts[3],
                        '段落': para_idx + 1
                    })
    
    print(f"从Word提取了 {len(data)} 条记录")
    return data

def compare_data(pdf_data, word_data):
    """比对PDF和Word数据"""
    print("\n开始比对数据...")
    
    # 转换为DataFrame便于比较
    pdf_df = pd.DataFrame(pdf_data)
    word_df = pd.DataFrame(word_data)
    
    print(f"PDF数据行数: {len(pdf_df)}")
    print(f"Word数据行数: {len(word_df)}")
    
    if len(pdf_df) == 0:
        print("❌ PDF数据为空!")
        return False
    
    if len(word_df) == 0:
        print("❌ Word数据为空!")
        return False
    
    # 按学号排序
    pdf_df = pdf_df.sort_values('学号').reset_index(drop=True)
    word_df = word_df.sort_values('学号').reset_index(drop=True)
    
    # 找出不匹配的记录
    mismatches = []
    all_student_ids = set(pdf_df['学号'].tolist() + word_df['学号'].tolist())
    
    for student_id in sorted(all_student_ids):
        pdf_row = pdf_df[pdf_df['学号'] == student_id]
        word_row = word_df[word_df['学号'] == student_id]
        
        if pdf_row.empty:
            mismatches.append(f"学号 {student_id}: 在PDF中未找到")
        elif word_row.empty:
            mismatches.append(f"学号 {student_id}: 在Word中未找到")
        else:
            pdf_grade = pdf_row.iloc[0]['智育成绩']
            word_grade = word_row.iloc[0]['智育成绩']
            
            if str(pdf_grade) != str(word_grade):
                mismatches.append(f"学号 {student_id}: 智育成绩不匹配 (PDF: {pdf_grade}, Word: {word_grade})")
    
    # 输出结果
    if not mismatches:
        print("✅ 所有数据完全一致!")
        print(f"共比对了 {len(all_student_ids)} 个学生的记录")
        return True
    else:
        print(f"❌ 发现 {len(mismatches)} 处不匹配:")
        for mismatch in mismatches[:20]:  # 只显示前20个不匹配项
            print(f"  - {mismatch}")
        
        if len(mismatches) > 20:
            print(f"  ... 还有 {len(mismatches) - 20} 个不匹配项")
        
        return False

def main():
    """主函数"""
    current_dir = Path(__file__).parent
    
    pdf_file = current_dir / "计算机学院（国家示范性软件学院）本科2023级计算机类2023-2024学年综合成绩公示.pdf"
    word_file = current_dir / "计算机学院（国家示范性软件学院）本科2023级计算机类2023-2024学年综合成绩公示.docx"
    
    if not pdf_file.exists():
        print(f"❌ PDF文件不存在: {pdf_file}")
        return
    
    if not word_file.exists():
        print(f"❌ Word文件不存在: {word_file}")
        return
    
    try:
        # 提取数据
        pdf_data = extract_pdf_data(pdf_file)
        word_data = extract_word_data(word_file)
        
        # 比对数据
        is_match = compare_data(pdf_data, word_data)
        
        # 保存详细比对结果到文件
        if pdf_data and word_data:
            pdf_df = pd.DataFrame(pdf_data)
            word_df = pd.DataFrame(word_data)
            
            # 保存提取的数据
            pdf_df.to_csv(current_dir / "pdf_extracted_data.csv", index=False, encoding='utf-8-sig')
            word_df.to_csv(current_dir / "word_extracted_data.csv", index=False, encoding='utf-8-sig')
            print(f"\n📊 详细数据已保存到:")
            print(f"  - pdf_extracted_data.csv")
            print(f"  - word_extracted_data.csv")
        
        if is_match:
            print("\n🎉 验证完成: PDF和Word文档数据完全一致!")
        else:
            print("\n⚠️  验证完成: 发现数据不一致，请查看上面的详细信息")
            
    except Exception as e:
        print(f"❌ 发生错误: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
